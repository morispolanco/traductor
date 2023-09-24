import streamlit as st
import requests
from docx import Document
from io import BytesIO
import json

# Cambiar el título en la pestaña del navegador
st.set_page_config(page_title="Traductor y Comparador de Texto", layout="centered")

# Clave API personalizada
secret_key = "9428f69325adc980cc9b9dc6a0f84a30a3eb86e74787792c581cc44e4c1adfae"

# Función para traducir texto al inglés
def translate_to_english(text):
    url = f"https://ai-translate.pro/api/{secret_key}/es-en"
    headers = {
        "Content-Type": "application/json"
    }
    data = {
        "text": text
    }

    response = requests.post(url, headers=headers, data=json.dumps(data))
    result = response.json()

    return result["result"]

# Función para traducir texto al español
def translate_to_spanish(text):
    url = f"https://ai-translate.pro/api/{secret_key}/en-es"
    headers = {
        "Content-Type": "application/json"
    }
    data = {
        "text": text
    }

    response = requests.post(url, headers=headers, data=json.dumps(data))
    result = response.json()

    return result["result"]

# Función para comparar textos
def compare_texts(original_text, translated_text):
    if original_text == translated_text:
        return "Los textos son iguales."
    else:
        return "Los textos son diferentes."

# Función para guardar la comparación en un documento Word
def save_comparison(original_text, translated_text):
    doc = Document()
    doc.add_heading("Comparación de Textos", level=1)
    doc.add_paragraph("Texto original:")
    doc.add_paragraph(original_text)
    doc.add_paragraph("Texto traducido:")
    doc.add_paragraph(translated_text)

    doc.save("comparison.docx")

def main():
    st.title("Traductor y Comparador de Texto")
    st.write("Cargue el documento Word en Español:")
    uploaded_file = st.file_uploader("Seleccione un archivo", type=["docx"])

    if uploaded_file is not None:
        doc = Document(uploaded_file)
        original_text = " ".join([p.text for p in doc.paragraphs])

        if st.button("Traducir al Inglés"):
            translated_text = translate_to_english(original_text)
            st.write("Texto traducido al Inglés:")
            st.write(translated_text)

            if st.button("Traducir al Español"):
                translated_back_text = translate_to_spanish(translated_text)
                st.write("Texto traducido al Español:")
                st.write(translated_back_text)

                comparison_result = compare_texts(original_text, translated_back_text)
                st.write("Comparación de textos:")
                st.write(comparison_result)

                if st.button("Guardar comparación"):
                    save_comparison(original_text, translated_back_text)
                    st.write("Comparación guardada con éxito.")

                    # Proporcionar enlace de descarga para el archivo de comparación
                    st.markdown("[Descargar comparación](comparison.docx)")

if __name__ == "__main__":
    main()
