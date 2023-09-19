import streamlit as st
import requests
from docx import Document
from io import BytesIO

# Cambiar el título en la pestaña del navegador
st.set_page_config(page_title="AITranslate", layout="centered")

# URL base de la API de AI Translate
BASE_URL = "https://ai-translate.pro/api"

# Función para traducir texto
def translate_text(text, lang_from, lang_to, secret_key):
    url = f"{BASE_URL}/{secret_key}/{lang_from}-{lang_to}"
    headers = {'Content-Type': 'application/json'}
    data = {"text": text}
    response = requests.post(url, headers=headers, json=data)
    if response.status_code == 200:
        result = response.json()["result"]
        available_chars = response.json()["available_chars"]
        return result, available_chars
    else:
        return None, None


# Título de la aplicación
st.title("AITranslate")

# Agregar título y texto en la parte superior
st.markdown("## La mejor traducción automática del mundo")
st.markdown("Las redes neuronales de AITranslate son capaces de captar hasta los más mínimos matices y reproducirlos en la traducción a diferencia de cualquier otro servicio. Para evaluar la calidad de nuestros modelos de traducción automática, realizamos regularmente pruebas a ciegas. En las pruebas a ciegas, los traductores profesionales seleccionan la traducción más precisa sin saber qué empresa la produjo. AITranslate supera a la competencia por un factor de 3:1.")

# Cargar archivo DOCX
uploaded_file = st.file_uploader("Cargar archivo DOCX", type=["docx"])

# Selección de idiomas
lang_from = st.selectbox("Seleccione el idioma de origen:", ["en", "es"])
lang_to = st.selectbox("Seleccione el idioma de destino:", ["en", "es"])

# Botón para traducir
if st.button("Traducir"):
    # Obtener la clave API desde los secretos
    secret_key = st.secrets["aitranslate"]["api_key"]

    if secret_key and uploaded_file is not None:
        # Leer el contenido del archivo DOCX
        docx = Document(uploaded_file)
        text = "\n".join([paragraph.text for paragraph in docx.paragraphs])

        translation, available_chars = translate_text(text, lang_from, lang_to, secret_key)
        if translation:
            # Crear un nuevo documento DOCX con la traducción
            translated_docx = Document()
            translated_docx.add_paragraph(translation)

            # Guardar el documento DOCX en un objeto BytesIO
            docx_buffer = BytesIO()
            translated_docx.save(docx_buffer)
            docx_buffer.seek(0)

            # Descargar el archivo DOCX
            st.download_button("Descargar traducción", data=docx_buffer, file_name="traduccion.docx")

            st.success("La traducción se ha guardado en el archivo 'traduccion.docx'")
            st.info(f"Caracteres disponibles: {available_chars}")
        else:
            st.error("Error al traducir el texto. Verifique su clave API o intente nuevamente.")
    else:
        st.error("Por favor, cargue un archivo DOCX.")
